"""
Extração de texto dos formatos suportados: PDF, DOCX, TXT, HTML, MD, CSV, JSON.

Cada carregador devolve ``(texto, metadados)``. As bibliotecas específicas são
opcionais: quando ausentes, o erro informa exatamente o que instalar.
"""

from __future__ import annotations

import csv
import io
import json
import logging
import re
from pathlib import Path
from typing import Any, Callable

from ...core.exceptions import ArquivoInvalido

logger = logging.getLogger(__name__)

# Limite de segurança para arquivos-texto lidos inteiros em memória.
MAX_CARACTERES = 20_000_000


# ---------------------------------------------------------------------- TXT/MD
def carregar_texto(caminho: Path) -> tuple[str, dict[str, Any]]:
    """Lê arquivos de texto puro, tolerando codificações diferentes de UTF-8."""
    dados = caminho.read_bytes()
    for codificacao in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            texto = dados.decode(codificacao)
            return texto[:MAX_CARACTERES], {"encoding": codificacao}
        except UnicodeDecodeError:
            continue
    # Último recurso: substitui os bytes inválidos em vez de falhar.
    return dados.decode("utf-8", errors="replace")[:MAX_CARACTERES], {
        "encoding": "utf-8-replace"
    }


# -------------------------------------------------------------------------- PDF
def carregar_pdf(caminho: Path) -> tuple[str, dict[str, Any]]:
    """Extrai o texto de um PDF página a página."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ArquivoInvalido(
            "Leitura de PDF requer a biblioteca 'pypdf'. Execute: pip install pypdf"
        ) from exc

    leitor = PdfReader(str(caminho))
    paginas: list[str] = []
    for numero, pagina in enumerate(leitor.pages, start=1):
        try:
            conteudo = pagina.extract_text() or ""
        except Exception as exc:
            logger.warning("Falha ao extrair página %d de %s: %s", numero, caminho, exc)
            conteudo = ""
        # Marcador de página: permite citar a origem exata na resposta do RAG.
        paginas.append(f"[[página {numero}]]\n{conteudo}")

    texto = "\n\n".join(paginas)
    metadados: dict[str, Any] = {"pages": len(leitor.pages)}
    if leitor.metadata:
        metadados.update(
            {
                "title": leitor.metadata.get("/Title"),
                "author": leitor.metadata.get("/Author"),
                "producer": leitor.metadata.get("/Producer"),
            }
        )

    # PDF de imagens escaneadas não tem camada de texto: sugere OCR.
    if len(texto.strip()) < 40 * len(leitor.pages):
        metadados["ocr_recomendado"] = True

    return texto, metadados


# ------------------------------------------------------------------------- DOCX
def carregar_docx(caminho: Path) -> tuple[str, dict[str, Any]]:
    """Extrai parágrafos e tabelas de um documento Word."""
    try:
        import docx
    except ImportError as exc:
        raise ArquivoInvalido(
            "Leitura de DOCX requer 'python-docx'. Execute: pip install python-docx"
        ) from exc

    documento = docx.Document(str(caminho))
    partes = [p.text for p in documento.paragraphs if p.text.strip()]

    for tabela in documento.tables:
        for linha in tabela.rows:
            celulas = [c.text.strip() for c in linha.cells]
            if any(celulas):
                partes.append(" | ".join(celulas))

    return "\n".join(partes), {
        "paragraphs": len(documento.paragraphs),
        "tables": len(documento.tables),
    }


# ------------------------------------------------------------------------- HTML
def carregar_html(caminho: Path) -> tuple[str, dict[str, Any]]:
    """Converte HTML em texto, descartando script/style/nav."""
    bruto, meta = carregar_texto(caminho)
    try:
        from bs4 import BeautifulSoup

        sopa = BeautifulSoup(bruto, "html.parser")
        for tag in sopa(["script", "style", "noscript", "nav", "footer", "svg"]):
            tag.decompose()
        titulo = sopa.title.string.strip() if sopa.title and sopa.title.string else None
        texto = sopa.get_text(separator="\n")
        meta["title"] = titulo
    except ImportError:
        # Fallback sem dependência: remoção de tags por expressão regular.
        logger.info("beautifulsoup4 ausente; usando extração simplificada de HTML.")
        sem_blocos = re.sub(
            r"<(script|style)[^>]*>.*?</\1>", " ", bruto, flags=re.S | re.I
        )
        texto = re.sub(r"<[^>]+>", " ", sem_blocos)
        texto = re.sub(r"&nbsp;?", " ", texto)

    return _limpar_espacos(texto), meta


# -------------------------------------------------------------------------- CSV
def carregar_csv(caminho: Path) -> tuple[str, dict[str, Any]]:
    """Converte um CSV em linhas legíveis 'coluna: valor'."""
    bruto, meta = carregar_texto(caminho)
    try:
        dialeto = csv.Sniffer().sniff(bruto[:8192])
    except csv.Error:
        dialeto = csv.excel  # separador padrão quando a detecção falha

    leitor = csv.reader(io.StringIO(bruto), dialeto)
    linhas = list(leitor)
    if not linhas:
        return "", {"rows": 0}

    cabecalho = linhas[0]
    partes = []
    for linha in linhas[1:]:
        pares = [
            f"{cabecalho[i] if i < len(cabecalho) else f'col{i}'}: {valor}"
            for i, valor in enumerate(linha)
            if valor.strip()
        ]
        if pares:
            partes.append("; ".join(pares))

    meta.update({"rows": len(linhas) - 1, "columns": cabecalho})
    return "\n".join(partes), meta


# ------------------------------------------------------------------------- JSON
def carregar_json(caminho: Path) -> tuple[str, dict[str, Any]]:
    """Achata um JSON em linhas 'caminho.da.chave: valor'."""
    bruto, meta = carregar_texto(caminho)
    try:
        dados = json.loads(bruto)
    except json.JSONDecodeError as exc:
        raise ArquivoInvalido(f"JSON inválido: {exc}") from exc

    linhas: list[str] = []
    _achatar(dados, "", linhas)
    meta["entries"] = len(linhas)
    return "\n".join(linhas), meta


def _achatar(valor: Any, prefixo: str, saida: list[str], profundidade: int = 0) -> None:
    """Percorre estruturas aninhadas produzindo linhas planas."""
    if profundidade > 12:  # proteção contra estruturas patologicamente aninhadas
        return
    if isinstance(valor, dict):
        for chave, item in valor.items():
            _achatar(item, f"{prefixo}.{chave}" if prefixo else str(chave), saida, profundidade + 1)
    elif isinstance(valor, list):
        for indice, item in enumerate(valor):
            _achatar(item, f"{prefixo}[{indice}]", saida, profundidade + 1)
    elif valor is not None and str(valor).strip():
        saida.append(f"{prefixo}: {valor}")


# --------------------------------------------------------------------- registro
CARREGADORES: dict[str, Callable[[Path], tuple[str, dict[str, Any]]]] = {
    ".txt": carregar_texto,
    ".md": carregar_texto,
    ".markdown": carregar_texto,
    ".pdf": carregar_pdf,
    ".docx": carregar_docx,
    ".html": carregar_html,
    ".htm": carregar_html,
    ".csv": carregar_csv,
    ".json": carregar_json,
}

EXTENSOES_SUPORTADAS = tuple(CARREGADORES)


def extrair(caminho: str | Path) -> tuple[str, dict[str, Any]]:
    """
    Extrai o texto de um arquivo pelo tipo, escolhendo o carregador adequado.

    Lança ``ArquivoInvalido`` para extensões sem suporte ou conteúdo vazio.
    """
    caminho = Path(caminho)
    if not caminho.exists():
        raise ArquivoInvalido(f"Arquivo não encontrado: {caminho.name}")

    carregador = CARREGADORES.get(caminho.suffix.lower())
    if carregador is None:
        raise ArquivoInvalido(
            f"Formato '{caminho.suffix}' não suportado. "
            f"Aceitos: {', '.join(sorted(EXTENSOES_SUPORTADAS))}"
        )

    texto, metadados = carregador(caminho)
    texto = _limpar_espacos(texto)

    if not texto.strip():
        dica = (
            " O documento parece ser digitalizado — tente o OCR em Documentos."
            if metadados.get("ocr_recomendado")
            else ""
        )
        raise ArquivoInvalido(f"Nenhum texto extraído de '{caminho.name}'.{dica}")

    metadados["chars"] = len(texto)
    return texto, metadados


def _limpar_espacos(texto: str) -> str:
    """Normaliza quebras de linha e remove espaçamento excessivo."""
    texto = texto.replace("\r\n", "\n").replace("\r", "\n")
    texto = re.sub(r"[ \t]+", " ", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()
